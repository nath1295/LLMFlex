from ..VectorDBs.base_vectordb import list_vectordbs
from ..Embeddings.base_embeddings import BaseEmbeddingsToolkit
from ..TextSplitters.base_text_splitter import BaseTextSplitter
from ..VectorDBs.base_vectordb import BaseVectorDatabase
from ..Models.Cores.base_core import BaseLLM
from ..Rankers.base_ranker import BaseRanker
from ..Schemas.documents import RankResult, Document
import os
from typing import List, Type, Optional, Callable, Literal, Tuple

def knowledge_base_dir() -> str:
    """Directory to store knowlege base.

    Returns:
        str: Directory to store knowlege base.
    """
    from ..utils import get_config
    kb_dir = os.path.join(get_config()['package_home'], 'knowledge_base')
    os.makedirs(kb_dir, exist_ok=True)
    return kb_dir

def list_knowledge_base() -> List[str]:
    """List the existing knowledge base.

    Returns:
        List[str]: List of the existing knowledge bases.
    """
    import re
    re_kb = re.compile('kb_\d+')
    dirs = os.listdir(knowledge_base_dir())
    dirs = list(filter(lambda x: re_kb.match(x), dirs))
    return dirs

def get_new_kb_id() -> str:
    """Get a new id for a new knowledge base.

    Returns:
        str: The new kb_id.
    """
    ids = list_knowledge_base()
    ids = list(map(lambda x: int(x.removeprefix('kb_')), ids))
    if len(ids) == 0:
        return 'kb_0'
    max_id = max(ids)
    new = max_id + 1
    for i in range(max_id):
        if i not in ids:
            new = i
            break
    return f'kb_{new}'

def load_markdown(file_dir: str) -> List[Document]:
    """Load a markdown file as list of documents for the knowledge base to add.

    Args:
        file_dir (str): Full directory of the markdown file.

    Returns:
        List[Document]: List of documents from the markdown file.
    """
    file_dir = os.path.abspath(file_dir)
    filename = os.path.basename(file_dir)
    with open(file_dir, 'r') as f:
        text = f.read()
    hash_start = text.startswith('#')
    chunks = filter(lambda x: x != '', text.split('\n#'))
    chunks = list(map(lambda x: ('#' + x).strip(), chunks))
    if hash_start:
        chunks[0] = chunks[0][1:]
    docs = list(map(lambda x: Document(index=x, metadata=dict(filename=filename, file_dir=str(file_dir))), chunks))
    return docs

def load_docx(file_dir: str) -> List[Document]:
    """Load a docx file as list of documents for the knowledge base to add.

    Args:
        file_dir (str): Full directory of the docx file.

    Returns:
        List[Document]: List of documents from the docx file.
    """
    from docx import Document as DocX
    file_dir = os.path.abspath(file_dir)
    filename = os.path.basename(file_dir)
    file = DocX(file_dir)
    chunks = filter(lambda x: x.text.strip() != '', file.paragraphs)
    chunks = map(lambda x: x.text.strip(), chunks)
    docs = list(map(lambda x: Document(index=x, metadata=dict(filename=filename, file_dir=str(file_dir))), chunks))
    return docs

def load_pdf(file_dir: str) -> List[Document]:
    """Load a pdf file as list of documents for the knowledge base to add.

    Args:
        file_dir (str): Full directory of the pdf file.

    Returns:
        List[Document]: List of documents from the pdf file.
    """
    import fitz
    file_dir = os.path.abspath(file_dir)
    filename = os.path.basename(file_dir)
    with fitz.open(file_dir) as doc:
        pages = []
        for page in doc:
            pages.append(page.get_text())

    docs = []
    for i, page in enumerate(pages):
        chunks = map(lambda x: x.strip(), page.split('\n\n'))
        chunks = filter(lambda x: x != '', chunks)
        chunks = list(map(lambda x: Document(index=x, metadata=dict(filename=filename, file_dir=str(file_dir), page=i)), chunks))
        docs.extend(chunks)
    return docs

def load_file(file_dir: str, filetype: Literal['auto', 'markdown', 'docx', 'pdf'] = 'auto') -> List[Document]:
    """Load a text-based file as list of docments.

    Args:
        file_dir (str): Full directory of the file.
        filetype (Literal[&#39;auto&#39;, &#39;markdown&#39;, &#39;docx&#39;, &#39;pdf&#39;], optional): The type of file to be loaded. If auto is set, it will be determined by the suffix of the file. Defaults to 'auto'.

    Returns:
        List[Document]: List of documents from the pdf file.
    """
    suffix = file_dir.split('.')[-1].lower()
    suf_map = dict(md='markdown', docx='docx', pdf='pdf')
    fn_map = dict(markdown=load_markdown, docx=load_docx, pdf=load_pdf)
    if filetype == 'auto':
        filetype = suf_map.get(suffix, 'markdown')
    return fn_map[filetype](file_dir=file_dir)
     

class KnowledgeBase:
    """Class to store any text as knowledge for querying.
    """
    def __init__(self, kb_id: str, embeddings: Type[BaseEmbeddingsToolkit], llm: Optional[BaseLLM], 
                 ranker: Optional[BaseRanker] = None,
                 text_splitter: Optional[BaseTextSplitter] = None,
                 ts_lang_model: str = 'en_core_web_sm',
                 chunk_size: int = 400,
                 chunk_overlap: int = 40) -> None:
        """Initialise the knowlege base.

        Args:
            kb_id (str): A unique identifier for the knowledge base starting with "kb_".
            embeddings (Type[BaseEmbeddingsToolkit]): Embeddings toolkit for the vector database.
            llm (Optional[BaseLLM]): LLM for counting tokens. If not given, the embedding model tokenizer will be used to count tokens.
            ranker (Optional[BaseRanker], optional): Reranker to rerank sementic search results. Defaults to None.
            text_splitter (Optional[BaseTextSplitter], optional): Text splitter to split documents. If None is given, it will be created with the token counting function. Defaults to None.
            ts_lang_model (str, optional): Text splitter language model to use if text_splitter is not provided. Defaults to 'en_core_web_sm'.
            chunk_size (int, optional): Chunk size of the text splitter if text_splitter is not provided. Defaults to 400.
            chunk_overlap (int, optional): Chunk overlap of the text splitter if text_splitter is not provided. Defaults to 40.
        """
        from ..TextSplitters.sentence_token_text_splitter import SentenceTokenTextSplitter
        from ..VectorDBs.faiss_vectordb import FaissVectorDatabase
        from ..Rankers.flashrank_ranker import FlashrankRanker
        self._kb_dir = os.path.join(knowledge_base_dir(), kb_id)
        self._embeddings = embeddings
        self._count_fn = self.embeddings.tokenizer.get_num_tokens if llm is None else llm.get_num_tokens
        self._text_splitter = SentenceTokenTextSplitter(
            count_token_fn=self.count_fn,
            language_model=ts_lang_model,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        ) if text_splitter is None else text_splitter
        self._ranker = FlashrankRanker() if ranker is None else ranker
        if kb_id in list_knowledge_base():
            self._vdb = FaissVectorDatabase.from_exist(embeddings=embeddings, name=kb_id, vectordb_dir=knowledge_base_dir(), text_splitter=self.text_splitter)
        else:
            import re
            if not re.match('kb_\d+', kb_id):
                raise ValueError(f'"{kb_id}" is not a valid kb_id, it must start with "kb_", followed by an integer.')
            self._vdb = FaissVectorDatabase.from_documents(embeddings=embeddings, docs=[], name=kb_id, vectordb_dir=knowledge_base_dir(), text_splitter=self.text_splitter)

    
    @property
    def embeddings(self) -> BaseEmbeddingsToolkit:
        """Embeddings toolkit for the vector database.

        Returns:
            BaseEmbeddingsToolkit: Embeddings toolkit for the vector database.
        """
        return self._embeddings
    
    @property
    def text_splitter(self) -> BaseTextSplitter:
        """Text splitter for the knowledge base.

        Returns:
            BaseTextSplitter: Text splitter for the knowledge base.
        """
        return self._text_splitter
    
    @property
    def vector_db(self) -> BaseVectorDatabase:
        """Vector database for the knowledge base.

        Returns:
            BaseVectorDatabase: Vector database for the knowledge base.
        """
        return self._vdb
    
    @property
    def knowledge_base_dir(self) -> str:
        """Directory for the knowledge base.

        Returns:
            str: Directory for the knowledge base.
        """
        if not os.path.exists(self._kb_dir):
            os.makedirs(self._kb_dir)
        return self._kb_dir
    
    @property
    def kb_id(self) -> str:
        """Knowledge base id.

        Returns:
            str: Knowledge base id.
        """
        return os.path.basename(self.knowledge_base_dir)

    @property
    def count_fn(self) -> Callable[[str], int]:
        """Function to count number of tokens in a string.

        Returns:
            Callable[[str], int]: Function to count number of tokens in a string.
        """
        return self._count_fn
    
    @property
    def ranker(self) -> BaseRanker:
        """Reranker for search results.

        Returns:
            BaseRanker: Reranker for search results.
        """
        return self._ranker

    @property
    def files(self) -> List[Tuple[str, str]]:
        """List of files and their respective directories.
        Returns:
            List[Tuple[str, str]]: List of files and their respective directories.
        """
        docs = self.vector_db.data.values()
        combos = list(set(map(lambda x: (x.metadata['filename'], x.metadata['file_dir']), docs)))
        return combos
    
    def search(self, query: str, top_k = 3, token_limit: Optional[int] = None, 
               fetch_k: int = 30,
               count_fn: Optional[Callable[[str], int]] = None,
               relevance_score_threshold: float = 0.8) -> List[RankResult]:
        """Searching for related information from the knowledge base.

        Args:
            query (str): Search query.
            top_k (int, optional): Maximum number of result. If token_limit is not None, token_limit will be used instead. Defaults to 3.
            token_limit (Optional[int], optional): Maximum number of tokens for the search results. Defaults to None.
            fetch_k (int, optional): Number of results to fetch from the vector database before reranking. Defaults to 30.
            count_fn (Optional[Callable[[str], int]], optional): Function to count the number of tokens if token_limit is not None. If None is given, the count_fn from the knowledge base class will be used. Defaults to None.
            relevance_score_threshold (float, optional): Minumum score for the reranking. Defaults to 0.8.

        Returns:
            List[RankResult]: List of search results.
        """
        count_fn = self.count_fn if count_fn is None else count_fn
        init_result = self.vector_db.search(query=query, top_k=fetch_k, index_only=False)
        if token_limit is None:
            result = self.ranker.rerank(query=query, elements=init_result, top_k=top_k)
        else: 
            rank_result = self.ranker.rerank(query=query, elements=init_result, top_k=len(init_result))
            result  =[]
            token_count = 0
            for res in rank_result:
                res_count = count_fn(res.index)
                if token_count + res_count <= token_limit:
                    token_count += res_count
                    result.append(res)
                else:
                    break
        result = list(filter(lambda x: x.rank_score >= relevance_score_threshold, result))
        return result
    
    def add_documents(self, docs: List[Document], mode: Literal['update', 'append'] = 'update') -> None:
        """Adding documents into the knowledge base. In the metadata of the file, it should contain at least filename and file_dir.

        Args:
            docs (List[Document]): List of documents to add.
            mode (Literal['update', 'append'], optional): Way of adding documents. Either updating/add the files or append on existing files. Defaults to 'update'.
        """
        val_docs = list(filter(lambda x: (('filename' in x.metadata.keys()) & ('file_dir' in x.metadata.keys())), docs))
        if len(val_docs) < len(docs):
            raise ValueError('"filename" and "file_dir" must exist in the metadata of all the documents.')
        if len(docs) != 0:
            if mode == 'update':
                combos = list(set(map(lambda x: (x.metadata['filename'], x.metadata['file_dir']), docs)))
                def filter_fn(doc: Document) -> bool:
                    combo = (doc.metadata['filename'], doc.metadata['file_dir'])
                    return combo in combos
                self.vector_db.delete_by_metadata(filter_fn=filter_fn)
            self.vector_db.add_documents(docs=docs, split_text=True, text_splitter=self.text_splitter)
                

    def clear(self) -> None:
        """Clear the entire knowledge base. Use it with caution.
        """
        self.vector_db.clear()

    

            


    


    

    
